import pickle
import uvicorn
import io
import os
import zipfile
import re
import nltk
import time
import uuid
import gc
from docx import Document
from shutil import make_archive, rmtree
from fastapi import FastAPI
from fastapi.middleware.cors import CORSMiddleware
from pydantic import BaseModel
from sklearn.model_selection import train_test_split
from sklearn.svm import LinearSVC
from sklearn.feature_extraction.text import CountVectorizer
from sklearn.pipeline import Pipeline
from sklearn.feature_extraction.text import TfidfTransformer
from sklearn.metrics import classification_report


class TrainDocumentsModel(BaseModel):
    path: str


class OrganizeFoldersModel(BaseModel):
    model_name: str
    organize_folder_path: str


class DeleteModelModel(BaseModel):
    model_id: str


app = FastAPI()
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

models_folder_name = 'local_models'
results_folder_name = 'results'

WPT = nltk.WordPunctTokenizer()
stop_word_list = ['acaba', 'ama', 'aslında', 'az', 'bazı', 'belki', 'biri', 'birkaç', 'birşey', 'biz', 'bu', 'çok',
                  'çünkü', 'da', 'daha', 'de', 'defa', 'diye', 'eğer', 'en', 'gibi', 'hem', 'hep', 'hepsi', 'her',
                  'hiç', 'için', 'ile', 'ise', 'kez', 'ki', 'kim', 'mı', 'mu', 'mü', 'nasıl', 'ne', 'neden', 'nerde',
                  'nerede', 'nereye', 'niçin', 'niye', 'o', 'sanki', 'şey', 'siz', 'şu', 'tüm', 've', 'veya', 'ya',
                  'yani']


def preprocess_doc(single_doc):
    # remove numbers
    number_pattern = re.compile('\d+')
    single_doc = number_pattern.sub('', single_doc)

    # remove punctuation and underscore
    punc_pattern = re.compile('[\W_]+')
    single_doc = punc_pattern.sub(' ', single_doc)

    # make lowercase
    single_doc = single_doc.lower()

    # remove trailing spaces
    single_doc = single_doc.strip()

    # remove multiple spaces
    single_doc = single_doc.replace('\s+', ' ')

    # remove stop words
    tokens = WPT.tokenize(single_doc)
    filtered_tokens = [token for token in tokens if token not in stop_word_list]
    single_doc = ' '.join(filtered_tokens)

    return single_doc


@app.post("/train_documents_with_path/")
async def train_documents_with_path(params: TrainDocumentsModel):
    start_time = time.time()

    text_list = []
    category_list = []
    zf = zipfile.ZipFile(params.path, "r")

    for file_name in zf.namelist():
        if not file_name.endswith("/"):
            doc = Document(io.BytesIO(zf.read(file_name)))
            text = ''.join([paragraph.text for paragraph in doc.paragraphs])
            text = preprocess_doc(text)
            text_list.append(text)

            category = file_name.split('/')[0]
            category_list.append(category)

    x_train, x_test, y_train, y_test = train_test_split(text_list, category_list, test_size=0.2)

    del zf
    del text_list
    del category_list
    gc.collect()

    svc = Pipeline([('vect', CountVectorizer()),
                    ('tfidf', TfidfTransformer()),
                    ('clf', LinearSVC())])
    svc.fit(x_train, y_train)
    y_prediction = svc.predict(x_test)
    prediction_report = classification_report(y_test, y_prediction)
    saved_model_file_name = f'{str(uuid.uuid4())}.pkl'

    del x_train
    del x_test
    del y_train
    del y_test
    gc.collect()

    if not os.path.exists(models_folder_name):
        os.makedirs(models_folder_name)
    pickle.dump(svc, open(f"{models_folder_name}/{saved_model_file_name}", 'wb'))

    return {
        "prediction_report": prediction_report,
        "elapsed_time": time.time() - start_time,
        "trained_model_name": saved_model_file_name
    }


@app.post("/organize_folders_with_path/")
async def organize_folders_with_path(params: OrganizeFoldersModel):
    model_file = open(f"{models_folder_name}/{params.model_name}", 'rb')
    model = pickle.load(model_file)

    results_file_name = str(uuid.uuid4())

    if not os.path.exists(results_folder_name):
        os.makedirs(results_folder_name)

    os.mkdir(f'{results_folder_name}/{results_file_name}')

    zf = zipfile.ZipFile(params.organize_folder_path, "r")
    for file_name in zf.namelist():
        doc = Document(io.BytesIO(zf.read(file_name)))
        text = ''.join([paragraph.text for paragraph in doc.paragraphs])
        text = preprocess_doc(text)
        prediction = model.predict([text])[0].strip()

        if not os.path.exists(f'{results_folder_name}/{results_file_name}/{prediction}'):
            os.makedirs(f'{results_folder_name}/{results_file_name}/{prediction}')

        doc.save(f'{results_folder_name}/{results_file_name}/{prediction}/{file_name}')

    make_archive(f'{results_folder_name}/{results_file_name}', "zip", f'{results_folder_name}/{results_file_name}')
    rmtree(f'{results_folder_name}/{results_file_name}', ignore_errors=True)
    return {
        "file_path": os.path.abspath(f'{results_folder_name}/{results_file_name}.zip')
    }


@app.post("/delete_model_file/")
async def delete_model_file(params: DeleteModelModel):
    os.remove(f"{models_folder_name}/{params.model_id}")
    return {
        "result": True
    }


if __name__ == "__main__":
    uvicorn.run(app, host="0.0.0.0", port=8000)
